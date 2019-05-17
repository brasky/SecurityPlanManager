from docx import *
from .models import Control, Implementation, Team
from django.db.models import Q
import re
from django.http import HttpResponse
from time import sleep
from django.contrib.staticfiles.storage import staticfiles_storage
import os 
from django.conf import settings
from .helper import get_control_parts
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
import sys
from django.db.models.query import QuerySet


namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main', 'w14': 'http://schemas.microsoft.com/office/word/2010/wordml' }


def generate_docx_ssp(baseline):
    template = __package__ + '\static\\fedramp_templates\\' + baseline + '.docx'
    document = Document(os.path.join(settings.BASE_DIR, template))
    previous_matched = []
    control_tables = {}
    control_to_implementation = {}
    for table in document.tables:
        try:   
            table_title = table.cell(0, 0).text
            table_title_column_two = table.cell(0, 1).text
            if "Control Summary Information" in table_title_column_two or "Control Enhancement Summary Information" in table_title_column_two:
                control_parent = table_title
                if "(" not in control_parent:
                    control = control_parent + "("
                    matching_controls = Control.objects.filter(Q(number__contains=control) & ~Q(number__contains=' '))
                    if not matching_controls:
                        control = control_parent
                        matching_controls = Control.objects.filter(Q(number__contains=control) & ~Q(number__contains=' '))
                else:
                    control = control_parent
                    matching_controls = Control.objects.filter(Q(number__contains=control))
                previous_matched = matching_controls
                for matched_control in matching_controls:
                    try:
                        matching_implementation = Implementation.objects.get(control=matched_control)
                        control_to_implementation[matched_control.number] = matching_implementation
                    except Implementation.MultipleObjectsReturned as e:
                        matching_implementation_group = Implementation.objects.filter(control=matched_control)
                        matching_implementation = matching_implementation_group[0]
                        control_to_implementation[matched_control.number] = matching_implementation_group
                        # print(matching_implementation)
                    
                    rows = len(table.rows)
                    # print(matched_control.number)
                    for row in range(1,rows):
                        cell_text = table.cell(row, 0).text
                        if "Responsible Role" in cell_text:
                            if matching_implementation.responsible_role not in cell_text:
                                table.cell(row,0).text = cell_text + " " + matching_implementation.responsible_role
                        elif "Parameter" in cell_text:
                            try:
                                if matched_control.number + ":" in cell_text:
                                    table.cell(row, 0).text = table.cell(row, 0).text + matching_implementation.parameter
                                elif '-1:' in cell_text and matched_control.number.replace(' ', '') in cell_text:
                                    sub_param = matching_implementation.parameter.split('2:')[0]
                                    table.cell(row, 0).text = table.cell(row, 0).text + ' ' + sub_param.replace('1:', '').strip()
                                elif '-2:' in cell_text and matched_control.number.replace(' ', '') in cell_text:
                                    sub_param = matching_implementation.parameter.split('3:')[0].split('2:')[1]
                                    table.cell(row, 0).text = table.cell(row, 0).text + sub_param.strip()
                                elif '-3:' in cell_text and matched_control.number.replace(' ', '') in cell_text:
                                    sub_param = matching_implementation.parameter.split(':4')[0].split('3:')[1]
                                    table.cell(row, 0).text = table.cell(row, 0).text + sub_param.strip()
                                elif '-4' in cell_text and matched_control.number.replace(' ', '') in cell_text:
                                    sub_param = matching_implementation.parameter.split('4:')[1]
                                    table.cell(row, 0).text = table.cell(row, 0).text + sub_param.strip()
                            except IndexError:
                                table.cell(row, 0).text = table.cell(row, 0).text + matching_implementation.parameter
                            
                        elif "Implementation" in cell_text:
                            implementation_status = matching_implementation.get_implementation_status_display()
                            for paragraph in table.cell(row, 0).paragraphs:
                                
                                p = paragraph._element
                                if len(p.xpath('.//w:t')) > 1:
                                    checkbox_implementation_status_text = p.xpath('.//w:t')[1].text.strip()
                                    if implementation_status.lower() == checkbox_implementation_status_text.lower():
                                        add_check_in_paragraph(p)
                        elif "Control Origination" in cell_text:
                            control_originations = matching_implementation.control_origination.all()

                            for control_origination_object in control_originations:
                                control_origination = control_origination_object.get_source_display()
                                for paragraph in table.cell(row, 0).paragraphs:
                                    p = paragraph._element
                                    if len(p.xpath('.//w:t')) > 1:
                                        checkbox_control_origination_status_text = p.xpath('.//w:t')[1].text.strip()
                                        if control_origination.lower() in checkbox_control_origination_status_text.lower():
                                            checkbox = p.find('.//w14:checkbox', namespace)
                                            if checkbox is not None:
                                                checkbox[0].set("{http://schemas.microsoft.com/office/word/2010/wordml}val", "1")
                                                p.xpath('.//w:t')[0].text = u'☒'
                                        elif "Inherited" in control_origination and "Inherited" in checkbox_control_origination_status_text:
                                            checkbox = p.find('.//w14:checkbox', namespace)
                                            if checkbox is not None:
                                                checkbox[0].set("{http://schemas.microsoft.com/office/word/2010/wordml}val", "1")
                                                p.xpath('.//w:t')[0].text = u'☒'                                            

            elif "solution" in table_title:
                matching_controls = previous_matched
                for matched_control in matching_controls:
                    control_parts = get_control_parts(matched_control.number)
                    # print("getting implementation")
                    implementation = control_to_implementation[matched_control.number]
                    # print(implementation)
                    if isinstance(implementation, QuerySet):
                        for imp in implementation:
                            # print("more than one imp")
                            add_implementation_to_table(table, imp, control_parts)
                    else:
                        # print("one or fewer imp")
                        add_implementation_to_table(table, implementation, control_parts)
        except Exception as e:
            print('Exception:')
            print(table_title)
            print(e)
            e = sys.exc_info()[0]
            print(e)


    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=high-ssp.docx'
    document.save(response)
    return response

def add_check_in_paragraph(paragraph):
    checkbox = paragraph.find('.//w14:checkbox', namespace)
    if checkbox is not None:
        checkbox[0].set("{http://schemas.microsoft.com/office/word/2010/wordml}val", "1")
        paragraph.xpath('.//w:t')[0].text = u'☒'
            
        
def add_implementation_to_table(table, implementation_object, control_parts):
    implementation_details = implementation_object.solution
    teams = implementation_object.teams.all()
    teams = ", ".join(map(str, teams))
    teams += ":"
    customer_responsibility = implementation_object.customer_responsibility
    if  not control_parts['part_letter'] and not control_parts['part_num']: #no letter, no part, no enhancement
        if customer_responsibility:
            if "Customer Responsibility:" not in customer_responsibility:
                table.cell(0, 1).text += "\n" + "Customer Responsibility:" + "\n" + customer_responsibility + "\n" + teams + "\n" + implementation_details
            else:
                table.cell(0, 1).text += "\n" + customer_responsibility + "\n" + teams + "\n" + implementation_details
        else:
            table.cell(0, 1).text += "\n" + teams + "\n" + implementation_details

    elif not control_parts['part_num']:#letter, no part, no enhancement
        
        if customer_responsibility:
            if "Customer Responsibility:" not in customer_responsibility:
                table.cell(control_parts['part_letter'], 1).text = (table.cell(control_parts['part_letter'], 1).text
                                                                + "Customer Responsibility:" + "\n"
                                                                + customer_responsibility + "\n"
                                                                + teams + "\n" + implementation_details + "\n")
            else:
                table.cell(control_parts['part_letter'], 1).text = (table.cell(control_parts['part_letter'], 1).text
                                                                    + customer_responsibility + "\n"
                                                                     + teams + "\n" + implementation_details + "\n")
        else:
            table.cell(control_parts['part_letter'], 1).text = (table.cell(control_parts['part_letter'], 1).text
                                                               + teams + "\n" + implementation_details + "\n")

    else:#letter, part, no enhancement
        if customer_responsibility:
            if "Customer Responsibility:" not in customer_responsibility:
                table.cell(control_parts['part_letter'], 1).text = (table.cell(control_parts['part_letter'], 1).text
                                                                + "Part " + control_parts['part_num'] + ":" + "\n"
                                                                + "Customer Responsibility:" +  "\n"
                                                                + customer_responsibility + "\n"
                                                                + teams + "\n" + implementation_details + "\n")
            else:
                table.cell(control_parts['part_letter'], 1).text = (table.cell(control_parts['part_letter'], 1).text
                                                                + "Part " + control_parts['part_num'] + ":" + "\n"
                                                                + customer_responsibility + "\n"
                                                                + teams + "\n" + implementation_details + "\n")
        else:
            table.cell(control_parts['part_letter'], 1).text = (table.cell(control_parts['part_letter'], 1).text
                                                               + "Part " + control_parts['part_num'] + ":" + "\n"
                                                               + teams + "\n" + implementation_details + "\n")
    